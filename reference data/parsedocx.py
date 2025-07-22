from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import pandas as pd
from typing import List, Tuple
import math
import numpy as np

def extract_color_df(docx_path, table_index):
    document = Document(docx_path)
    specific_table_colors_list = []

    if table_index < 0:
        print("Error: table_index must be a non-negative integer.")
        return pd.DataFrame()

    if table_index >= len(document.tables):
        print(f"Error: Table with index {table_index} does not exist. Document has only {len(document.tables)} tables.")
        return pd.DataFrame()

    target_table = document.tables[table_index]

    for row in target_table.rows:
        row_data = []
        for cell in row.cells:
            shading_color = None
            try:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = tc_pr.find(
                    'w:shd',
                    {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                )
                if shd is not None:
                    shading_color = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                else:
                    shading_color = cell.text.strip()
            except Exception as e:
                print(f"Error processing cell: {e}")
            row_data.append(shading_color)
        specific_table_colors_list.append(row_data)

    df = pd.DataFrame(specific_table_colors_list)

    return df

def extract_table_headers(docx_path, table_index):
    """
    Extracts header names (text from the first row) of a specific table
    (by index) in a DOCX file.

    Args:
        docx_path (str): The path to the DOCX file.
        table_index (int): The 0-based index of the table to extract headers from.

    Returns:
        list: A list of strings representing the header names from the first row.
              Returns an empty list if the table_index is out of range,
              no tables are found, or the specified table has no rows.
    """
    document = Document(docx_path)
    headers = []

    if table_index < 0:
        print("Error: table_index must be a non-negative integer.")
        return []

    if table_index >= len(document.tables):
        print(f"Error: Table with index {table_index} does not exist. Document has only {len(document.tables)} tables.")
        return []

    target_table = document.tables[table_index]

    if len(target_table.rows) > 0:
        header_row = target_table.rows[0]
        for cell in header_row.cells:
            headers.append(cell.text.strip())
    else:
        print(f"Warning: Table at index {table_index} has no rows. Cannot extract headers.")
        return []

    return headers


doc = "mou\\LIST OF MOU PARTNER Collections 2025-26.docx"
legend = extract_color_df(doc, 0)

legend1 = legend.iloc[:, :2].rename(columns={0:"color", 1:"status"}).replace(r'^\s*$', pd.NA, regex=True).dropna()
legend2 = legend.iloc[:, 2:4].rename(columns={2:"color", 3:"status"}).replace(r'^\s*$', pd.NA, regex=True).dropna()
legend3 = legend.iloc[:, 4:6].rename(columns={4:"color", 5:"status"}).replace(r'^\s*$', pd.NA, regex=True).dropna()

legends = pd.concat([legend1,legend2, legend3], ignore_index=True)
legends[['colorname', 'status']] = legends['status'].str.split(':', expand=True)

#print(legends)


head_current = extract_table_headers(doc,1)
current = extract_color_df(doc,1)
current.columns = head_current

todrop = ['Ready to Initiate', 'Financial codes confirmed', 'Collection Completed']
current = current.drop(columns=todrop)
current = current.rename(columns={"":"Status" , "Status": "Status_comment"})

#print(current)
#print(legends)

dict_legends = legends.set_index('color')['status'].to_dict()  

# Convert hex color to RGB tuple
def hex_to_rgb(hex_code: str) -> Tuple[int, int, int]:
    hex_code = hex_code.strip("#")
    return tuple(int(hex_code[i:i+2], 16) for i in (0, 2, 4))

# Compute Euclidean distance between two RGB colors
def color_distance(rgb1: Tuple[int, int, int], rgb2: Tuple[int, int, int]) -> float:
    return math.sqrt(sum((a - b) ** 2 for a, b in zip(rgb1, rgb2)))

# Find the closest match from the legend for each color in the dataset
def find_closest_colors(dataset_colors: List[str]) -> List[Tuple[str, str, float]]:
    legend_rgb = {k: hex_to_rgb(k) for k in dict_legends}
    results = []

    for color in dataset_colors:
        color_rgb = hex_to_rgb(color)
        closest_match = None
        min_distance = float('inf')
        for legend_hex, legend_rgb_val in legend_rgb.items():
            dist = color_distance(color_rgb, legend_rgb_val)
            if dist < min_distance:
                min_distance = dist
                closest_match = legend_hex
        results.append((color, closest_match, dict_legends[closest_match]))

    return results


matches = find_closest_colors(current['Status'])

# Convert matches to DataFrame
matches_df = pd.DataFrame(matches, columns=['Status', 'ClosestColor', 'LegendStatus'])

# Join by index to avoid duplication
updateddf = current.copy()
updateddf[['ClosestColor', 'LegendStatus']] = matches_df[['ClosestColor', 'LegendStatus']].values

filtered_df = updateddf[updateddf['Status'] != '000000']

#print(filtered_df)

# Match up with partner data 

file_moucorrected = "partnerdata.csv"

dfn = pd.read_csv(file_moucorrected)

dfn['gc_orgID'] = dfn['gc_orgID'].apply(lambda x: str(int(x)) if pd.notnull(x) else '')
dfn = dfn[dfn['gc_orgID'] != '']
mou = pd.DataFrame(dfn , columns =['gc_orgID', 'partner'])

#print(mou)

identifer_df = filtered_df.merge(mou, left_on="Org", right_on="partner")
identifer_df = identifer_df.drop(columns=["Org", "Status", "Level"])
identifer_df = identifer_df.rename(columns={"ClosestColor":"Status" , "partner": "Org"})

#print(identifer_df.columns)

identifer_df['Contribution'] = (
    identifer_df['Contribution']
    .str.replace("\n", "")
    .str.replace("In-Kind", "-inf")
    .str.strip()
    .str.replace(",", "")
    .astype(float)
)

identifer_df['Next steps'] = identifer_df['Next steps'].str.replace("FFFFFF","")

print(identifer_df)

# Look up partner levels
partnerlevels = pd.read_csv('mou\\mou_levels.csv')

# Function to find level
def get_level(amount):
    for _, row in partnerlevels.iterrows():
        if row['Minimum'] <= amount <= row['Maximum']:
            return row['Level']
        if amount == float('-inf') :
            return "In-Kind"
    return None

# Apply function to assign levels
identifer_df['Level'] = identifer_df['Contribution'].apply(get_level)

desired_order = ['Status', "LegendStatus",'Org', 'Level', 'Contribution', 'Status_comment', 'Next steps']

output = identifer_df[desired_order]

print(output['Level'])

# Save as a CSV
output.to_csv('moucollections.csv', index=False, encoding='utf-8-sig')


# Extract Potential Partner List
head_potential = extract_table_headers(doc,2)
potential = extract_color_df(doc,2)


document_partners = Document(doc)
